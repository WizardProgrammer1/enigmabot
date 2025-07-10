import whisper
import torch
import logging
import os
from src.config import Config
from docx import Document
from datetime import timedelta
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle
from src.file_handler import FileHandler
import noisereduce as nr
from scipy.io import wavfile
import tempfile
from src.transcription_providers import TranscriptionFactory

print("=== THIS IS THE ACTIVE transcriber.py ===")

class Transcriber:
    def __init__(self, provider_name: str = None):
        # Создаем провайдер транскрипции
        self.provider_name = provider_name or Config.TRANSCRIPTION_PROVIDER
        logging.info(f"Инициализация провайдера транскрипции: {self.provider_name}")
        
        try:
            self.provider = TranscriptionFactory.create_provider(self.provider_name)
            logging.info(f"Провайдер {self.provider_name} успешно инициализирован")
        except Exception as e:
            logging.error(f"Ошибка инициализации провайдера {self.provider_name}: {e}")
            # Fallback на локальный Whisper
            logging.info("Переключение на локальный Whisper")
            self.provider = TranscriptionFactory.create_provider('whisper_local')
            self.provider_name = 'whisper_local'

        self.file_handler = FileHandler()

    def _denoise_audio(self, file_path: str) -> str:
        """Применяет шумоподавление к аудиофайлу и возвращает путь к временному очищенному файлу."""
        rate, data = wavfile.read(file_path)
        # Если стерео, берем только один канал
        if len(data.shape) > 1:
            data = data[:, 0]
        reduced_noise = nr.reduce_noise(y=data, sr=rate)
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
            wavfile.write(tmp.name, rate, reduced_noise.astype(data.dtype))
            return tmp.name

    def transcribe(self, file_path: str, language: str = 'ru') -> str:
        print("=== TRANSCRIBE CALLED ===", file_path, language)
        try:
            print("=== INSIDE TRANSCRIBE TRY ===")
            # --- DENOISE AUDIO ---
            denoised_path = self._denoise_audio(file_path)
            
            # Используем выбранный провайдер
            if self.provider_name == 'whisper_local' and self.get_duration(file_path) > Config.CHUNK_LENGTH * 2:
                # Для локального Whisper используем чанки для длинных файлов
                result = self._transcribe_long_file_whisper(denoised_path, language)
            else:
                # Для остальных провайдеров используем прямой вызов
                result = self.provider.transcribe(denoised_path, language)
            
            # Удаляем временный файл
            if os.path.exists(denoised_path):
                os.remove(denoised_path)
            return result
        except Exception as e:
            print(f"[ERROR] В transcribe: {e}")
            raise

    def _transcribe_short_file_whisper(self, file_path: str, language: str = 'ru') -> str:
        print(f"=== _transcribe_short_file_whisper CALLED === {file_path}")
        options = Config.WHISPER_OPTIONS.copy()
        if language and language != 'other':
            options['language'] = language
        options['task'] = 'transcribe'
        print(f"[WHISPER] SHORT: file={file_path}, options={options}")
        if hasattr(self.provider, "model"):
            result = self.provider.model.transcribe(file_path, **options)
            return str(result['text'])
        else:
            return str(self.provider.transcribe(file_path, language))

    def _transcribe_long_file_whisper(self, file_path: str, language: str = 'ru') -> str:
        print(f"=== _transcribe_long_file_whisper CALLED === {file_path}")
        import ffmpeg
        import tempfile
        print("Используем транскрипцию по частям для длинного файла")
        duration = self.get_duration(file_path)
        chunk_length = Config.CHUNK_LENGTH
        chunks = []
        for start in range(0, int(duration), chunk_length):
            end = min(start + chunk_length, int(duration))
            chunks.append((start, end))
        print(f"Файл разбит на {len(chunks)} частей")
        results = []
        for i, (start, end) in enumerate(chunks):
            print(f"Обрабатываем часть {i+1}/{len(chunks)} ({start}-{end}с)")
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
                tmp_chunk_path = tmp.name
            try:
                (
                    ffmpeg
                    .input(file_path, ss=start, t=end-start)
                    .output(tmp_chunk_path, format='wav', acodec='pcm_s16le', ac=1, ar='16000', loglevel='error')
                    .overwrite_output()
                    .run()
                )
                print(f"[WHISPER] LONG: chunk={tmp_chunk_path}")
                if hasattr(self.provider, "model"):
                    text = str(self.provider.model.transcribe(tmp_chunk_path, **Config.WHISPER_OPTIONS)['text']).strip()
                else:
                    text = str(self.provider.transcribe(tmp_chunk_path, language)).strip()
                results.append(text)
            except Exception as e:
                print(f"Ошибка при обработке чанка {i+1}: {e}")
            finally:
                if os.path.exists(tmp_chunk_path):
                    os.remove(tmp_chunk_path)
        full_text = ' '.join(results)
        print(f"Транскрипция завершена. Длина текста: {len(full_text)} символов")
        return full_text

    def get_duration(self, file_path: str) -> float:
        """Получает длительность файла в секундах"""
        return self.file_handler.get_duration(file_path)

    def get_text_with_timestamps(self, file_path: str, language: str = 'ru') -> str:
        """Возвращает текст с таймкодами для файла."""
        options = Config.WHISPER_OPTIONS.copy()
        if language and language != 'other':
            options['language'] = language
        options['task'] = 'transcribe'
        print(f"[WHISPER] TIMESTAMPS: file={file_path}, options={options}")
        # Универсальная поддержка таймкодов
        if hasattr(self.provider, "transcribe_with_timestamps"):
            result = self.provider.transcribe_with_timestamps(file_path, language)
            segments = result.get("segments", [])
            if not segments:
                return result.get("text", "")
        elif hasattr(self.provider, "model"):
            result = self.provider.model.transcribe(file_path, **options)
            segments = result.get("segments", [])
        else:
            text = self.provider.transcribe(file_path, language)
            return text
        lines = []
        for seg in segments:
            if isinstance(seg, dict):
                start = int(seg.get("start", 0))
                start_str = str(timedelta(seconds=start))
                text_val = seg.get("text", "")
                lines.append(f"[{start_str}] {text_val.strip()}")
            else:
                lines.append(str(seg))
        return "\n".join(lines)

    def get_docx_with_timestamps(self, file_path: str, language: str = 'ru', out_path: str = None) -> str:
        """Генерирует docx-файл с текстом и таймкодами, возвращает путь к файлу."""
        options = Config.WHISPER_OPTIONS.copy()
        if language and language != 'other':
            options['language'] = language
        options['task'] = 'transcribe'
        print(f"[WHISPER] DOCX: file={file_path}, options={options}")
        from docx import Document
        doc = Document()
        # Универсальная поддержка таймкодов
        if hasattr(self.provider, "transcribe_with_timestamps"):
            result = self.provider.transcribe_with_timestamps(file_path, language)
            segments = result.get("segments", [])
            if not segments:
                doc.add_heading('Транскрипция', 0)
                doc.add_paragraph(result.get("text", ""))
                if not out_path:
                    import tempfile
                    fd, out_path_tmp = tempfile.mkstemp(suffix='.docx')
                    os.close(fd)
                    out_path = out_path_tmp
                doc.save(out_path)
                return out_path
        elif hasattr(self.provider, "model"):
            result = self.provider.model.transcribe(file_path, **options)
            segments = result.get("segments", [])
        else:
            text = self.provider.transcribe(file_path, language)
            doc.add_heading('Транскрипция', 0)
            doc.add_paragraph(text)
            if not out_path:
                import tempfile
                fd, out_path_tmp = tempfile.mkstemp(suffix='.docx')
                os.close(fd)
                out_path = out_path_tmp
            doc.save(out_path)
            return out_path
        doc.add_heading('Транскрипция с таймкодами', 0)
        for seg in segments:
            if isinstance(seg, dict):
                start = int(seg.get("start", 0))
                start_str = str(timedelta(seconds=start))
                text_val = seg.get("text", "")
                doc.add_paragraph(f"[{start_str}] {text_val.strip()}")
            else:
                doc.add_paragraph(str(seg))
        if not out_path:
            import tempfile
            fd, out_path_tmp = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            out_path = out_path_tmp
        doc.save(out_path)
        return out_path

    def upload_docx_to_gdrive(self, docx_path: str, filename: str = None) -> str:
        """
        Загружает docx-файл на Google Drive и возвращает публичную ссылку на Google Docs.
        Имя документа в Google Docs теперь всегда берется из filename (title).
        """
        creds = None
        if os.path.exists('token.json'):
            creds = Credentials.from_authorized_user_file('token.json', ['https://www.googleapis.com/auth/drive.file'])
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                from google_auth_oauthlib.flow import InstalledAppFlow
                if not Config.GOOGLE_CLIENT_ID or not Config.GOOGLE_CLIENT_SECRET:
                    raise ValueError("GOOGLE_CLIENT_ID и GOOGLE_CLIENT_SECRET должны быть установлены в переменных окружения")
                flow = InstalledAppFlow.from_client_config({
                    "installed": {
                        "client_id": Config.GOOGLE_CLIENT_ID,
                        "client_secret": Config.GOOGLE_CLIENT_SECRET,
                        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                        "token_uri": "https://oauth2.googleapis.com/token",
                        "redirect_uris": ["http://localhost"]
                    }
                }, ['https://www.googleapis.com/auth/drive.file'])
                creds = flow.run_local_server(port=0)
            with open('token.json', 'w') as token:
                token.write(creds.to_json())

        service = build('drive', 'v3', credentials=creds)
        # Имя документа в Google Docs — всегда из filename
        file_metadata = {
            'name': filename or os.path.basename(docx_path),
            'mimeType': 'application/vnd.google-apps.document'
        }
        media = MediaFileUpload(docx_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        file_id = file.get('id')
        # Делаем файл публичным
        service.permissions().create(fileId=file_id, body={'type': 'anyone', 'role': 'reader'}).execute()
        # Получаем ссылку на Google Docs
        link = f'https://docs.google.com/document/d/{file_id}/edit'
        return link

    async def transcribe_long_file_with_progress(self, file_path: str, options: dict, progress_callback=None) -> str:
        """Асинхронная транскрипция длинных файлов по частям с прогрессом."""
        import ffmpeg
        import tempfile
        import asyncio
        logging.info("Используем транскрипцию по частям для длинного файла (async)")
        duration = self.get_duration(file_path)
        chunk_length = Config.CHUNK_LENGTH
        chunks = []
        for start in range(0, int(duration), chunk_length):
            end = min(start + chunk_length, int(duration))
            chunks.append((start, end))
        logging.info(f"Файл разбит на {len(chunks)} частей")
        results = []
        for i, (start, end) in enumerate(chunks):
            logging.info(f"Обрабатываем часть {i+1}/{len(chunks)} ({start}-{end}с)")
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
                tmp_chunk_path = tmp.name
            try:
                (
                    ffmpeg
                    .input(file_path, ss=start, t=end-start)
                    .output(tmp_chunk_path, format='wav', acodec='pcm_s16le', ac=1, ar='16000', loglevel='error')
                    .overwrite_output()
                    .run()
                )
                if hasattr(self.provider, "model"):
                    text = str(self.provider.model.transcribe(tmp_chunk_path, **options)['text']).strip()
                else:
                    text = str(self.provider.transcribe(tmp_chunk_path, options.get('language', 'ru'))).strip()
                results.append(text)
            except Exception as e:
                logging.error(f"Ошибка при обработке чанка {i+1}: {e}")
            finally:
                if os.path.exists(tmp_chunk_path):
                    os.remove(tmp_chunk_path)
            # Прогресс
            if progress_callback:
                percent = int((i+1)/len(chunks)*100)
                await progress_callback(percent)
            await asyncio.sleep(0)  # дать другим таскам поработать
        full_text = self.format_txt(' '.join(results))
        logging.info(f"Транскрипция завершена. Длина текста: {len(full_text)} символов")
        return full_text

    async def transcribe_long_file_with_progress_async(self, file_path: str, options: dict, progress_callback=None) -> str:
        """
        Асинхронно выполняет транскрибацию длинного файла в отдельном потоке, чтобы не блокировать event loop.
        Прогресс-коллбек не будет вызываться асинхронно, но основной поток не блокируется.
        """
        import asyncio
        loop = asyncio.get_event_loop()
        # Исправлено: вызываем актуальный метод и передаем язык
        return await loop.run_in_executor(None, lambda: self._transcribe_long_file_whisper(file_path, options.get('language', 'ru')))

    def format_txt(self, text: str) -> str:
        """Красиво форматирует текст для txt-файла: абзацы, переносы, убирает лишние пробелы."""
        import re
        # Удаляем лишние пробелы
        text = re.sub(r' +', ' ', text)
        # Разбиваем на предложения и делаем абзацы
        sentences = re.split(r'(?<=[.!?]) +', text)
        paragraphs = []
        para = []
        for sent in sentences:
            para.append(sent)
            if len(' '.join(para)) > 300:
                paragraphs.append(' '.join(para))
                para = []
        if para:
            paragraphs.append(' '.join(para))
        return '\n\n'.join(paragraphs) 